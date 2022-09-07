from docx import Document
#from docx.shared import Inches
from docx import shared
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #对齐方式
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import time,os

def table_head_bg(table,coln,color):
	color_bg = locals()
	color_bg = {}
	for x in range(coln):
		color_bg[x] = parse_xml(r'<w:shd {} w:fill="{color}"/>'.format(nsdecls('w'),color=color))
		table.rows[0].cells[x]._tc.get_or_add_tcPr().append(color_bg[x])

def run(c,data1_result,baseinfo,inspection_data_result,hostdata):
	report_filename = '{taskid}_{host}_{port}_{time}.docx'.format(taskid=baseinfo['taskid'], task_detail_id=baseinfo['task_detail_id'], host=baseinfo['host'], port=baseinfo['port'], time=str(int(time.time()))) #要返回的巡检报告的名字
	report_filepath = os.path.abspath(baseinfo['report_dir'])
	report_file = str(report_filepath) + '/' + str(report_filename) #用于生成巡检报告的(绝对路径)

	document = Document()

	#设置全局设置. 比如字体
	#document.styles['Normal'].font.name = u'微软雅黑'
	document.styles['Normal'].font.name = u'微软雅黑'
	document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑') #设置中文字体使用字体2->宋体

	document.add_heading('MYSQL巡检报告', 0)
#	p = document.add_paragraph('脚本地址: http://www.github.com/ddcw/inspection')
#	p.add_run('bold').bold = True
#	p.add_run(' and some ')
#	p.add_run('italic.').italic = True

#	p2 = document.add_paragraph(baseinfo)
#	p3 = document.add_paragraph(data1_result)
#	p4 = document.add_paragraph(inspection_data_result)
#	p5 = document.add_paragraph(hostdata)

	document.add_heading('基础信息', level=1)
	#document.add_paragraph('Intense quote(测试样式)', style='Intense Quote')

	#基础信息的表
	table = document.add_table(rows=1, cols=2) #定义一个表
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = '对象'
	hdr_cells[1].text = '值'
	row_cells = table.add_row().cells
	row_cells[0].text = 'MYSQL地址'
	row_cells[1].text = baseinfo['host']
	row_cells2 = table.add_row().cells
	row_cells2[0].text = 'MYSQL端口'
	row_cells2[1].text = str(baseinfo['port'])
	row_cells3 = table.add_row().cells
	row_cells3[0].text = 'MYSQL版本'
	row_cells3[1].text = str(baseinfo['version'])
	row_cells4 = table.add_row().cells
	row_cells4[0].text = 'MYSQL使用内存'
	row_cells4[1].text = str(baseinfo['mem_total'])
	row_cells5 = table.add_row().cells
	row_cells5[0].text = '巡检时间'
	row_cells5[1].text = str(baseinfo['begin_time'])
	row_cells6 = table.add_row().cells
	row_cells6[0].text = '巡检成功项'
	row_cells6[1].text = str(baseinfo['inspection_success'])
	row_cells7 = table.add_row().cells
	row_cells7[0].text = '各阶段状态'
	row_cells7[1].text = "正常:{inspection_level_1} 提示:{inspection_level_2} 警告:{inspection_level_3} 严重:{inspection_level_4} 致命:{inspection_level_5}".format(inspection_level_1=baseinfo['inspection_level_1'], inspection_level_2=baseinfo['inspection_level_2'], inspection_level_3=baseinfo['inspection_level_3'], inspection_level_4=baseinfo['inspection_level_4'], inspection_level_5=baseinfo['inspection_level_5'])
	row_cells8 = table.add_row().cells
	row_cells8[0].text = '巡检得分(百分比)'
	row_cells8[1].text = str(baseinfo['score'])
	row_cells9 = table.add_row().cells
	row_cells9[0].text = '角色'
	row_cells9[1].text = str(baseinfo['role'])
	row_cells10 = table.add_row().cells
	row_cells10[0].text = '数据库运行时间'
	row_cells10[1].text = '{d} 天'.format(d=baseinfo['uptime'])

	#库表信息
	document.add_heading('数据库信息', level=1)
	db_table_table = document.add_table(rows=1, cols=4) #定义一个表
	db_table_table_head = db_table_table.rows[0].cells
	db_table_table_head[0].text = '库名'
	db_table_table_head[1].text = '数据大小'
	db_table_table_head[2].text = '索引大小'
	db_table_table_head[3].text = '默认字符集'
	for x in baseinfo['db_table']['data'].values:
		db_table_table_row = db_table_table.add_row().cells
		db_table_table_row[0].text = x[0]
		db_table_table_row[1].text = "{s}GB".format(s=round(int(x[1])/1024/1024/1024,2))
		db_table_table_row[2].text = "{s}GB".format(s=round(int(x[2])/1024/1024/1024,2))
		db_table_table_row[3].text = x[3]

	#主机信息(hostdata)
	if hostdata['status']:
		document.add_heading('主机相关信息', level=1)
		document.add_heading('基础信息', level=2)
		hosttable1 = document.add_table(rows=1, cols=2) #定义一个表
		hostinfo1 = hosttable1.rows[0].cells
		hostinfo1[0].text = '对象'
		hostinfo1[1].text = '值'
		hostinfo2 = hosttable1.add_row().cells
		hostinfo2[0].text = '主机名'
		hostinfo2[1].text = str(hostdata['data']['hostname']['stdout'])
		hostinfo3 = hosttable1.add_row().cells
		hostinfo3[0].text = '架构'
		hostinfo3[1].text = str(hostdata['data']['platform']['stdout'])
		hostinfo4 = hosttable1.add_row().cells
		hostinfo4[0].text = '内核'
		hostinfo4[1].text = str(hostdata['data']['kernel']['stdout'])
		hostinfo5 = hosttable1.add_row().cells
		hostinfo5[0].text = 'cpu数量'
		hostinfo5[1].text = '{socket} * {core} * {thread} = {total}'.format(socket=hostdata['data']['cpu_socket']['stdout'], core=hostdata['data']['cpu_core']['stdout'], thread=hostdata['data']['cpu_thread']['stdout'], total=int(hostdata['data']['cpu_socket']['stdout'])*int(hostdata['data']['cpu_core']['stdout'])*int(hostdata['data']['cpu_thread']['stdout']))
		hostinfo6 = hosttable1.add_row().cells
		hostinfo6[0].text = 'cpu型号'
		hostinfo6[1].text = str(hostdata['data']['cpu_name']['stdout'])
		hostinfo7 = hosttable1.add_row().cells
		hostinfo7[0].text = '总内存(GB)'
		hostinfo7[1].text = str(hostdata['data']['mem_total']['stdout'])
		hostinfo8 = hosttable1.add_row().cells
		hostinfo8[0].text = '可用内存(GB)'
		hostinfo8[1].text = str(hostdata['data']['mem_available']['stdout'])
		hostinfo9 = hosttable1.add_row().cells
		hostinfo9[0].text = '总SWAP'
		hostinfo9[1].text = '{swap}GB(swappiness={swappiness})'.format(swap=hostdata['data']['swap_total']['stdout'],swappiness=hostdata['data']['swappiness']['stdout'])
		hostinfo10 = hosttable1.add_row().cells
		hostinfo10[0].text = '负载'
		hostinfo10[1].text = str(hostdata['data']['loadavg']['stdout'])
		hostinfo11 = hosttable1.add_row().cells
		hostinfo11[0].text = '时区'
		hostinfo11[1].text = str(hostdata['data']['timezone']['stdout'])
		
		document.add_heading('数据库目录信息', level=2)
		hosttable2= document.add_table(rows=1, cols=8) #定义一个表
		hostinfo1 = hosttable2.rows[0].cells
		hostinfo1[0].text = '对象'
		hostinfo1[1].text = '文件系统'
		hostinfo1[2].text = '类型'
		hostinfo1[3].text = '总大小(GB)'
		hostinfo1[4].text = '使用大小(GB)'
		hostinfo1[5].text = '可用大小(GB)'
		hostinfo1[6].text = '使用百分比'
		hostinfo1[7].text = '挂载点'
		if hostdata['data']['dirstatus']:
			for x in ['datadir','redodir','binlogdir','tmpdir']:
				hostinfo_dir = hosttable2.add_row().cells
				hostinfo_dir[0].text = str(x)
				_tmp_n = 1
				for y in hostdata['data'][x]['stdout']:
					hostinfo_dir[_tmp_n].text = str(y)
					_tmp_n += 1
		

	#数据库引擎相关(baseinfo)
	document.add_heading('数据库引擎相关', level=1)
	dbengine_table = document.add_table(rows=1, cols=2) #定义一个表
	dbengine_row = dbengine_table.rows[0].cells
	dbengine_row[0].text = '对象'
	dbengine_row[1].text = '值(%)'
	for dbengine_row_n in ['key_buffer_read_hits','key_buffer_write_hits','query_cache_hits','innodb_buffer_read_hits','table_open_cache_hits']:
		dbengine_rows = dbengine_table.add_row().cells
		dbengine_rows[0].text = dbengine_row_n
		dbengine_rows[1].text = str(baseinfo[dbengine_row_n])
	dbengine_row1 = dbengine_table.add_row().cells
	dbengine_row1[0].text = 'INNODB内存使用率(总内存{total_mem}MB)'.format(total_mem=baseinfo['innodb_mem_total'])
	dbengine_row1[1].text = str(baseinfo['innodb_mem_used_p'])

	#数据库参数
	document.add_heading('数据库参数', level=1)
	db_table = document.add_table(rows=1, cols=4) #定义一个表
	db_table_head = db_table.rows[0].cells
	db_table_head[0].text = '参数'
	db_table_head[1].text = '值'
	db_table_head[2].text = '描述'
	db_table_head[3].text = '建议'
	for x in inspection_data_result:
		if inspection_data_result[x]['status'] == False or inspection_data_result[x]['type'] != 5 :
			continue
		db_table_row = db_table.add_row().cells
		db_table_row[0].text = x
		db_table_row[1].text = str(inspection_data_result[x]['data'].values[0][0])
		db_table_row[2].text = inspection_data_result[x]['des']
		if inspection_data_result[x]['score'] < inspection_data_result[x]['old_score']:
			db_table_row[3].text = inspection_data_result[x]['suggestion']
		

	#数据库状态(TODO)
	#document.add_heading('数据库状态', level=1)
	#p = document.add_paragraph('TO BE CONTINUE....')
	#TO BE CONTINUE (计划使用plt画图...)

	#巡检项汇总(不含参数)
	document.add_heading('巡检项汇总信息', level=1)
	table_inspection_sumary = document.add_table(rows=1, cols=7) #定义一个表
	row0 = table_inspection_sumary.rows[0].cells
	row0[0].text = '巡检项'
	row0[1].text = '描述'
	row0[2].text = '是否巡检成功'
	row0[3].text = '分类'
	row0[4].text = '得分'
	row0[5].text = '总分'
	row0[6].text = '状态'
	for x in inspection_data_result:
		if not inspection_data_result[x]['status']:
			continue
		inspection_sumary_row = table_inspection_sumary.add_row().cells
		inspection_sumary_row[0].text = str(x)
		inspection_sumary_row[1].text = str(inspection_data_result[x]['des'])
		#inspection_sumary_row[2].text = str(inspection_data_result[x]['status'])
		inspection_sumary_row[2].text = '巡检成功' if inspection_data_result[x]['status'] else '巡检失败'
		#inspection_sumary_row[3].text = str(inspection_data_result[x]['type'])
		inspection_sumary_row[3].text = c['OTHER']['inspection_type'][inspection_data_result[x]['type']][0]
		inspection_sumary_row[4].text = str(inspection_data_result[x]['score'])
		inspection_sumary_row[5].text = str(inspection_data_result[x]['old_score'])
		#inspection_sumary_row[6].text = str(inspection_data_result[x]['level'])
		inspection_sumary_row[6].text = c['OTHER']['default_name'][inspection_data_result[x]['level']][0]
		


	#巡检项详情
	document.add_heading('巡检项详情', level=1)

	#巡检详情
	for x in inspection_data_result:
		if inspection_data_result[x]['status'] == False or inspection_data_result[x]['type'] == 5 :
			continue
		#document.add_heading(x, level=2)
		document.add_heading(inspection_data_result[x]['des'], level=2)
		table2 = document.add_table(rows=1, cols=inspection_data_result[x]['data'].shape[1]) #定义一个表
		row1 = table2.rows[0].cells
		#表头
		row_title_n = 0
		for row_title in inspection_data_result[x]['data']:
			row1[row_title_n].text = str(row_title)
			row_title_n += 1

		#设置表头背景颜色
		table_head_bg(table2,inspection_data_result[x]['data'].shape[1],'DCDCDC')

		#数据
		for row_data in inspection_data_result[x]['data'].values:
			row_data_n = 0
			row_detail = table2.add_row().cells
			for col in row_data:
				row_detail[row_data_n].text = str(col)
				row_data_n += 1
		
	#一些图(又得引入matplotlib.pyplot了.....)	

	#巡检建议
	document.add_heading('巡检建议', level=1)
	suggestion_1 = document.add_table(rows=0, cols=1) #定义一个表
	#suggestion_info1 = suggestion_1.rows[0].cells
	#suggestion_info1[0].text = '建议'
	for x in baseinfo['inspection_suggestion']:
		suggestion_info2 = suggestion_1.add_row().cells
		suggestion_info2[0].text = str(x)


	#document.add_page_break()
	document.save(report_file)

	return {'type':'docx','status':True,'data':report_filename}
